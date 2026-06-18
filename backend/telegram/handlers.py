import os
import logging
import uuid
import tempfile
from telegram import Update
from telegram.ext import ContextTypes
from sqlalchemy.orm import Session
from backend.database.db import SessionLocal
from backend.database.models import Student, Document
from backend.services import global_services

# Try docx import
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Try pypdf import
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

logger = logging.getLogger(__name__)

def get_db():
    return SessionLocal()

async def handle_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = str(update.effective_chat.id)
    text = update.message.text.strip()
    
    # Check if there is a linking code
    # e.g., /start TUTOR-12345 or /link TUTOR-12345
    linking_code = None
    if text.startswith("/start ") and len(text) > 7:
        linking_code = text[7:].strip()
    elif text.startswith("/link ") and len(text) > 6:
        linking_code = text[6:].strip()
        
    db = get_db()
    try:
        # Check if already linked
        student = db.query(Student).filter(Student.telegram_id == chat_id).first()
        if student:
            await update.message.reply_text(
                f"Chào mừng quay trở lại, {student.name}! Tôi là AI Tutor của bạn. Hãy chat với tôi hoặc gửi tài liệu (PDF, Word, TXT) để học nhé!"
            )
            return

        if linking_code:
            # Attempt to link using code
            student = db.query(Student).filter(Student.linking_code == linking_code).first()
            if student:
                student.telegram_id = chat_id
                db.commit()
                await update.message.reply_text(
                    f"Liên kết thành công! Chào {student.name}, từ bây giờ bạn có thể tương tác với AI Tutor trực tiếp qua Telegram."
                )
            else:
                await update.message.reply_text(
                    "Mã liên kết không hợp lệ. Vui lòng kiểm tra lại mã trên giao diện Web."
                )
        else:
            # Create a new student profile directly from Telegram details
            name = update.effective_user.first_name or "Học sinh Telegram"
            new_code = f"TUTOR-{uuid.uuid4().hex[:6].upper()}"
            student = Student(
                telegram_id=chat_id,
                name=name,
                linking_code=new_code,
                learning_goals="Học tập tổng hợp",
                skill_level="Beginner"
            )
            db.add(student)
            db.commit()
            db.refresh(student)
            await update.message.reply_text(
                f"Xin chào {name}! Tôi đã tạo tài khoản AI Tutor cho bạn.\n\n"
                f"Mã liên kết Web của bạn là: `{new_code}`\n\n"
                f"Hãy chat hoặc gửi file cho tôi để bắt đầu học nhé!"
            )
    except Exception as e:
        logger.error(f"Error in handle_start: {e}")
        await update.message.reply_text("Có lỗi xảy ra khi bắt đầu. Vui lòng thử lại sau.")
    finally:
        db.close()

def chunk_text(text: str, max_length: int = 4000) -> list:
    chunks = []
    while len(text) > max_length:
        split_idx = text.rfind('\n', 0, max_length)
        if split_idx == -1:
            split_idx = text.rfind(' ', 0, max_length)
        if split_idx == -1:
            split_idx = max_length
        chunks.append(text[:split_idx])
        text = text[split_idx:].lstrip()
    if text:
        chunks.append(text)
    return chunks

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = str(update.effective_chat.id)
    user_text = update.message.text
    
    db = get_db()
    try:
        # Fetch student
        student = db.query(Student).filter(Student.telegram_id == chat_id).first()
        if not student:
            await update.message.reply_text("Vui lòng gõ /start để đăng ký hoặc liên kết tài khoản trước.")
            return

        # Show typing status
        await context.bot.send_chat_action(chat_id=chat_id, action="typing")
        
        # Call tutor service
        if global_services.tutor_service:
            reply, action = await global_services.tutor_service.handle_message(db, student.id, user_text)
            chunks = chunk_text(reply, max_length=4000)
            for chunk in chunks:
                try:
                    await update.message.reply_text(chunk, parse_mode="Markdown")
                except Exception as parse_err:
                    logger.warning(f"Failed to parse Markdown for Telegram reply chunk: {parse_err}. Falling back to plain text.")
                    await update.message.reply_text(chunk)
        else:
            await update.message.reply_text("Dịch vụ gia sư AI hiện chưa được khởi động.")
    except Exception as e:
        logger.error(f"Error in handle_message: {e}")
        await update.message.reply_text("Rất tiếc, tôi đang gặp lỗi kết nối. Hãy thử lại sau ít phút.")
    finally:
        db.close()

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = str(update.effective_chat.id)
    document = update.message.document
    filename = document.file_name
    
    db = get_db()
    try:
        # Fetch student
        student = db.query(Student).filter(Student.telegram_id == chat_id).first()
        if not student:
            await update.message.reply_text("Vui lòng gõ /start trước khi gửi tài liệu.")
            return
            
        file_ext = os.path.splitext(filename)[1].lower()
        if file_ext not in [".pdf", ".docx", ".txt"]:
            await update.message.reply_text("Tôi chỉ hỗ trợ các định dạng file PDF, DOCX và TXT.")
            return
            
        await update.message.reply_text(f"Đang tải và xử lý tài liệu '{filename}'...")
        
        # Download document
        telegram_file = await context.bot.get_file(document.file_id)
        
        # Save to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
            tmp_path = tmp_file.name
            await telegram_file.download_to_drive(tmp_path)
            
        # Parse content
        content = ""
        if file_ext == ".txt":
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        elif file_ext == ".pdf":
            if not PDF_AVAILABLE:
                await update.message.reply_text("Thư viện xử lý PDF chưa được cài đặt.")
                return
            reader = PdfReader(tmp_path)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content += text + "\n"
        elif file_ext == ".docx":
            if not DOCX_AVAILABLE:
                await update.message.reply_text("Thư viện xử lý DOCX chưa được cài đặt.")
                return
            doc = docx.Document(tmp_path)
            for para in doc.paragraphs:
                if para.text:
                    content += para.text + "\n"
                    
        # Clean up temporary file
        try:
            os.unlink(tmp_path)
        except Exception as e:
            logger.error(f"Error removing temp file: {e}")
            
        if not content.strip():
            await update.message.reply_text("Không tìm thấy nội dung văn bản trong tài liệu này.")
            return
            
        # Save document info to database
        storage_dir = "backend/storage/documents"
        os.makedirs(storage_dir, exist_ok=True)
        permanent_path = os.path.join(storage_dir, f"{uuid.uuid4().hex}{file_ext}")
        
        with open(permanent_path, "w", encoding="utf-8") as f:
            f.write(content)
            
        db_doc = Document(
            student_id=student.id,
            filename=filename,
            file_path=permanent_path,
            file_type=file_ext[1:]
        )
        db.add(db_doc)
        db.commit()
        db.refresh(db_doc)
        
        # Index document into ChromaDB (RAG)
        if global_services.rag_service:
            global_services.rag_service.add_document(student.id, db_doc.id, filename, content)
        
        await update.message.reply_text(
            f"Đã xử lý tài liệu '{filename}' thành công! Bây giờ bạn có thể đặt câu hỏi dựa trên nội dung tài liệu này."
        )
    except Exception as e:
        logger.error(f"Error in handle_document: {e}")
        await update.message.reply_text("Có lỗi xảy ra khi xử lý tài liệu của bạn.")
    finally:
        db.close()

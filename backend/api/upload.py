import os
import uuid
import tempfile
from fastapi import APIRouter, Depends, HTTPException, Cookie, UploadFile, File, Request
from sqlalchemy.orm import Session
from typing import Optional
from backend.database.db import get_db
from backend.database.models import Document

# Try docx import
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Try pypdf import
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Try PyMuPDF import
try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

router = APIRouter(prefix="/api/upload", tags=["upload"])

@router.post("")
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    student_id: Optional[str] = Cookie(None),
    db: Session = Depends(get_db)
):
    if not student_id:
        raise HTTPException(status_code=401, detail="Not logged in")
        
    student_id_int = int(student_id)
    filename = file.filename
    file_ext = os.path.splitext(filename)[1].lower()
    
    if file_ext not in [".pdf", ".docx", ".txt"]:
        raise HTTPException(status_code=400, detail="Only PDF, DOCX and TXT files are supported")
        
    try:
        import time
        start_time = time.time()
        total_token_usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
        
        content = ""
        # Write file content to a temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
            tmp_path = tmp_file.name
            file_data = await file.read()
            tmp_file.write(file_data)
            
        # Parse text based on type
        if file_ext == ".txt":
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        elif file_ext == ".pdf":
            if PYMUPDF_AVAILABLE:
                doc = fitz.open(tmp_path)
                llm_service = request.app.state.llm_service
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text = page.get_text()
                    images = page.get_images()
                    
                    has_images = len(images) > 0
                    has_drawings = False
                    if len(text.strip()) < 100:
                        try:
                            has_drawings = len(page.get_drawings()) > 0
                        except Exception:
                            pass
                    
                    if has_images or has_drawings:
                        # Page has images or graphics. Render to PNG.
                        try:
                            pix = page.get_pixmap(dpi=150)
                            img_bytes = pix.tobytes("png")
                            
                            prompt = (
                                f"Đây là trang {page_num + 1} của tài liệu học tập PDF tải lên. Trang này chứa hình ảnh, hình vẽ, sơ đồ hoặc đồ thị. "
                                "Hãy đọc kỹ, trích xuất toàn bộ văn bản và cung cấp mô tả chi tiết, đầy đủ nhất cho các hình ảnh, sơ đồ hoặc đồ thị "
                                "trên trang này bằng tiếng Việt để hỗ trợ việc dạy học của AI Tutor. "
                                "Đảm bảo giữ lại các công thức toán học và nội dung bài tập."
                            )
                            
                            page_description = await llm_service.analyze_page_image(img_bytes, prompt)
                            
                            usage = getattr(llm_service, "last_token_usage", None)
                            if usage:
                                total_token_usage["prompt_tokens"] += usage.get("prompt_tokens", 0)
                                total_token_usage["completion_tokens"] += usage.get("completion_tokens", 0)
                                total_token_usage["total_tokens"] += usage.get("total_tokens", 0)

                            if page_description:
                                content += f"\n--- Trang {page_num + 1} (Phân tích từ ảnh) ---\n{page_description}\n"
                            else:
                                content += f"\n--- Trang {page_num + 1} (Văn bản gốc) ---\n{text}\n"
                        except Exception as e:
                            content += f"\n--- Trang {page_num + 1} (Lỗi phân tích ảnh - Văn bản gốc) ---\n{text}\n"
                    else:
                        # Page only has text, extract directly
                        if text.strip():
                            content += f"\n--- Trang {page_num + 1} (Văn bản) ---\n{text}\n"
            elif PDF_AVAILABLE:
                reader = PdfReader(tmp_path)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
            else:
                raise HTTPException(status_code=500, detail="PDF parser dependency not installed")
        elif file_ext == ".docx":
            if not DOCX_AVAILABLE:
                raise HTTPException(status_code=500, detail="DOCX parser dependency not installed")
            doc = docx.Document(tmp_path)
            for para in doc.paragraphs:
                if para.text:
                    content += para.text + "\n"
                    
        # Clean up temp file
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
            
        if not content.strip():
            raise HTTPException(status_code=400, detail="No readable text found in the file")
            
        # Save file to permanent storage
        storage_dir = "backend/storage/documents"
        os.makedirs(storage_dir, exist_ok=True)
        permanent_path = os.path.join(storage_dir, f"{uuid.uuid4().hex}{file_ext}")
        
        with open(permanent_path, "w", encoding="utf-8") as f:
            f.write(content)
            
        # Save to database
        db_doc = Document(
            student_id=student_id_int,
            filename=filename,
            file_path=permanent_path,
            file_type=file_ext[1:]
        )
        db.add(db_doc)
        db.commit()
        db.refresh(db_doc)
        
        # Add to RAG collection
        rag_service = request.app.state.rag_service
        if rag_service:
            rag_service.add_document(student_id_int, db_doc.id, filename, content)
            
        end_time = time.time()
        exec_time = round(end_time - start_time, 2)
        
        return {
            "status": "success",
            "document": {
                "id": db_doc.id,
                "filename": db_doc.filename,
                "file_type": db_doc.file_type
            },
            "execution_time_seconds": exec_time,
            "token_usage": total_token_usage if total_token_usage["total_tokens"] > 0 else None
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/documents")
def get_documents(
    student_id: Optional[str] = Cookie(None),
    db: Session = Depends(get_db)
):
    if not student_id:
        raise HTTPException(status_code=401, detail="Not logged in")
        
    documents = db.query(Document).filter(Document.student_id == int(student_id)).all()
    return [
        {
            "id": doc.id,
            "filename": doc.filename,
            "file_type": doc.file_type,
            "created_at": doc.created_at.isoformat()
        }
        for doc in documents
    ]
